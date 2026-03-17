"""Tests for PandocConverter: verifies EPUB3 output structure and manifest."""

import hashlib
import io
import zipfile

import pytest


def _make_docx(tmp_path) -> tuple[str, bytes]:
    """Create a minimal DOCX fixture using python-docx and return (path, bytes)."""
    from docx import Document  # type: ignore[import]

    doc = Document()
    doc.add_heading("HITL Conversion Test", level=1)
    doc.add_paragraph("Hello, EPUB Conversion!")
    doc.add_paragraph("This is a minimal fixture document.")

    docx_path = tmp_path / "fixture.docx"
    doc.save(str(docx_path))
    return str(docx_path), docx_path.read_bytes()


def test_pandoc_converts_docx_to_epub(tmp_path):
    from src.converters.pandoc_converter import PandocConverter

    docx_path, docx_bytes = _make_docx(tmp_path)

    converter = PandocConverter()
    epub_bytes, manifest = converter.convert(docx_path)

    # Output must be a ZIP (EPUB is a ZIP archive)
    assert epub_bytes[:2] == b"PK", "EPUB output should start with ZIP magic bytes PK"

    # Must contain a 'mimetype' entry with the correct value
    with zipfile.ZipFile(io.BytesIO(epub_bytes)) as zf:
        assert "mimetype" in zf.namelist(), "EPUB ZIP must contain a 'mimetype' entry"
        mimetype = zf.read("mimetype").decode("utf-8").strip()
        assert mimetype == "application/epub+zip", (
            f"mimetype entry must equal 'application/epub+zip', got {mimetype!r}"
        )

    # Manifest must carry the correct SHA-256 of the original source
    expected_hash = hashlib.sha256(docx_bytes).hexdigest()
    assert manifest["sourceFileHash"] == expected_hash, (
        "manifest.sourceFileHash must be the SHA-256 of the source DOCX bytes"
    )
    assert manifest["sourceFormat"] == "docx"
    assert "convertedAt" in manifest
    assert isinstance(manifest["degradationNotices"], list)


def test_pandoc_converts_md_to_epub(tmp_path):
    from src.converters.pandoc_converter import PandocConverter

    md_path = tmp_path / "fixture.md"
    md_bytes = b"# Heading\n\nSome **bold** text and a paragraph.\n"
    md_path.write_bytes(md_bytes)

    epub_bytes, manifest = PandocConverter().convert(str(md_path))

    assert epub_bytes[:2] == b"PK"
    assert manifest["sourceFormat"] == "md"
    assert manifest["sourceFileHash"] == hashlib.sha256(md_bytes).hexdigest()

    with zipfile.ZipFile(io.BytesIO(epub_bytes)) as zf:
        assert zf.read("mimetype").strip() == b"application/epub+zip"
